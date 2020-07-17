import docx
import platform
from functions.functionLibrary import *
import openpyxl
from pandas import *
import PyPDF2
from bs4 import BeautifulSoup
from enum import Enum
from classes.clsReport import *

"""
Used Pandas to convert Excel Table to Python Dictionary
pandas API - https://pandas.pydata.org/pandas-docs/stable/user_guide/io.html
Alot less code with this implementation and more readable

Used an ENUM to type out Headers
"""

# start word document
my_document = docx.Document()


""" Used an Enum to print out the headers
"""
HEADING = Report.HEADING
OPERATIONS = Report.OPERATIONS
SALES = Report.SALES
MARKETING = Report.MARKETING
IT = Report.IT


"""Adds the paragraphs from the word file
"""
paragraphCount = 0
operationsWordDocFileName = 'text_files\operations.docx'
wordDocText = getText(operationsWordDocFileName)
wordDocTextPieces = wordDocText.split("\n")
my_document.add_heading(HEADING.value)
my_document.add_heading(OPERATIONS.value)
while paragraphCount < len(wordDocTextPieces):
    my_document.add_paragraph(wordDocTextPieces[paragraphCount])
    paragraphCount += 1


"""Adds the excel file table
"""
xls = ExcelFile('text_files\sales.xlsx')
excelData = xls.parse(xls.sheet_names[0])

excelDictObj = excelData.to_dict()

dfObj = pandas.DataFrame(excelDictObj)
my_document.add_heading(SALES.value)
tableObj = my_document.add_table(dfObj.shape[0]+1, dfObj.shape[1])


for value in range(dfObj.shape[-1]):
    tableObj.cell(0, value).text = dfObj.columns[value]


for item in range(dfObj.shape[0]):
    for text in range(dfObj.shape[-1]):
        tableObj.cell(item+1, text).text = str(dfObj.values[item, text])

"""Adds the text from the PDF file
"""
pdfFileObj = open('text_files\marketing.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
pageObj = pdfReader.getPage(0)

pdfTextObj = pageObj.extractText()

pdfTextObjPieces = pdfTextObj.split(".\n")

my_document.add_heading(MARKETING.value)

for wordBlock in pdfTextObjPieces:
    wordBlock = str(wordBlock).replace("\n", "")
    if "!" in wordBlock:
        wordBlock = wordBlock.split("!")
        counter = 0
        for sentenceBlock in wordBlock:

            if counter < 1:
                sentenceBlock = sentenceBlock + "!"
            else:
                sentenceBlock = sentenceBlock + "."

            if len(sentenceBlock) > 2:
                my_document.add_paragraph(sentenceBlock)

            counter += 1
    else:
        if len(wordBlock) > 2:
            wordBlock = wordBlock + ".\n"
            my_document.add_paragraph(wordBlock)


"""Adds the text from the HTML file
"""

reportText = open('text_files\IT.html', 'rb')
beautifulSoupText = BeautifulSoup(reportText, "html.parser")
htmlData = beautifulSoupText.findAll("p")

cleanHTML = []
my_document.add_heading(IT.value)
for htmlBlock in htmlData:
    htmlBlock = str(htmlBlock).replace("<p>", "")
    htmlBlock = str(htmlBlock).replace("</p>", "")
    cleanHTML.append(htmlBlock)

for paragraph in cleanHTML:
    my_document.add_paragraph(paragraph)
    
"""Save File
"""

my_document.save("text_files\ceo_report.docx")
